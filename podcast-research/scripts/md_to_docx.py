"""
Markdown to Word Converter
Converts Markdown files to Word documents (.docx) with full formatting support.

Requirements:
    python -m pip install markdown2 python-docx

Usage:
    python md_to_docx.py input.md output.docx
"""

import argparse
import re
import sys
from pathlib import Path

try:
    import markdown2
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError as e:
    print(f"Missing required library: {e}")
    print("Install with: python -m pip install markdown2 python-docx")
    sys.exit(1)


class MarkdownToWordConverter:
    """Convert Markdown to Word document with full formatting support."""

    def __init__(self, chinese_font="宋体", english_font="Arial", font_size=11):
        """
        Initialize the converter.

        Args:
            chinese_font: Font name for Chinese characters
            english_font: Font name for English characters
            font_size: Base font size in points
        """
        self.chinese_font = chinese_font
        self.english_font = english_font
        self.font_size = font_size
        self.doc = Document()

        # Set default styles
        self._setup_styles()

    def _setup_styles(self):
        """Configure document styles."""
        # Normal paragraph style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = self.english_font
        font.size = Pt(self.font_size)
        font.element.rPr.rFonts.set(qn('w:eastAsia'), self.chinese_font)

        # Heading styles
        for i in range(1, 7):
            heading_style = self.doc.styles[f'Heading {i}']
            heading_font = heading_style.font
            heading_font.name = self.english_font
            heading_font.size = Pt(16 - i)
            heading_font.bold = True
            heading_font.element.rPr.rFonts.set(qn('w:eastAsia'), self.chinese_font)

    def _set_run_font(self, run, bold=False, italic=False, code=False):
        """Set font properties for a text run."""
        run.font.name = self.english_font
        run.font.size = Pt(self.font_size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.element.rPr.rFonts.set(qn('w:eastAsia'), self.chinese_font)

        if code:
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B)

    def _add_page_break(self):
        """Add a page break."""
        self.doc.add_page_break()

    def parse_markdown_elements(self, md_text):
        """Parse Markdown into structured elements."""
        elements = []
        lines = md_text.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i]

            # Code block detection
            if line.strip().startswith('```'):
                code_lines = []
                lang = line.strip()[3:].strip() or 'text'
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                elements.append(('code', '\n'.join(code_lines), lang))
                i += 1
                continue

            # Horizontal rule
            if re.match(r'^[-*_]{3,}\s*$', line.strip()):
                elements.append(('hr',))
                i += 1
                continue

            # ATX style headings (# Heading)
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                content = heading_match.group(2)
                elements.append(('heading', content, level))
                i += 1
                continue

            # Setext style headings (underlined with === or ---)
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if next_line.startswith('===') and len(next_line) >= 3:
                    elements.append(('heading', line, 1))
                    i += 2
                    continue
                if next_line.startswith('---') and len(next_line) >= 3:
                    elements.append(('heading', line, 2))
                    i += 2
                    continue

            # Table detection (must come before blockquote since tables also use |)
            # Check if current line looks like a table header
            line_stripped = line.strip()
            if (line_stripped.startswith('|') and line_stripped.endswith('|') and
                line_stripped.count('|') >= 3):
                # Check if next line is a table separator
                if i + 1 < len(lines):
                    sep_line = lines[i + 1].strip()
                    # Table separator: must start/end with | and contain |- sequences
                    # Pattern like: |---|---:|:---:|
                    sep_ok = (sep_line.startswith('|') and sep_line.endswith('|') and
                             sep_line.count('|') >= 2 and
                             all('|' + cell + '|' in sep_line or
                                 re.match(r':?-+:?', cell.strip())
                                 for cell in sep_line.strip('|').split('|') if cell.strip()))
                    if sep_ok:
                        # It's a table! Parse header row
                        headers = [cell.strip() for cell in line_stripped.strip('|').split('|')]
                        i += 1  # Skip to separator line
                        i += 1  # Skip separator line
                        rows = []
                        # Collect all table rows (consecutive lines starting with |)
                        while i < len(lines):
                            row_line = lines[i].strip()
                            if row_line.startswith('|') and row_line.endswith('|'):
                                cells = [cell.strip() for cell in row_line.strip('|').split('|')]
                                rows.append(cells)
                                i += 1
                            else:
                                break
                        elements.append(('table', headers, rows))
                        continue

            # Blockquote
            if line.strip().startswith('>'):
                quote_lines = []
                while i < len(lines) and lines[i].strip().startswith('>'):
                    quote_lines.append(lines[i].strip()[1:].strip())
                    i += 1
                elements.append(('quote', '\n'.join(quote_lines)))
                continue

            # Unordered list
            if re.match(r'^[\s]*[-*+]\s+', line):
                list_items = []
                while i < len(lines) and re.match(r'^[\s]*[-*+]\s+', lines[i]):
                    list_items.append(re.sub(r'^[\s]*[-*+]\s+', '', lines[i]))
                    i += 1
                elements.append(('ulist', list_items))
                continue

            # Ordered list
            if re.match(r'^[\s]*\d+\.\s+', line):
                list_items = []
                while i < len(lines) and re.match(r'^[\s]*\d+\.\s+', lines[i]):
                    list_items.append(re.sub(r'^[\s]*\d+\.\s+', '', lines[i]))
                    i += 1
                elements.append(('olist', list_items))
                continue

            # Empty line
            if not line.strip():
                if elements and elements[-1][0] == 'paragraph':
                    elements[-1] = ('paragraph', elements[-1][1] + '\n')
                i += 1
                continue

            # Regular paragraph
            if line.strip():
                para_text = line
                i += 1
                # Don't merge lines that start with | (tables), headings, lists, quotes, or code blocks
                while i < len(lines) and lines[i].strip() and not re.match(r'^(#{1,6}|[-*+]|\d+\.|>|\|)', lines[i]):
                    para_text += '\n' + lines[i]
                    i += 1
                elements.append(('paragraph', para_text))
                continue

            i += 1

        return elements

    def _parse_inline_formatting(self, text):
        """Parse inline formatting (bold, italic, code, links, images)."""
        parts = []

        # Pattern for various inline elements
        patterns = [
            (r'!\[([^\]]*)\]\(([^\)]+)\)', 'image'),      # ![alt](url)
            (r'\[([^\]]+)\]\(([^\)]+)\)', 'link'),        # [text](url)
            (r'`([^`]+)`', 'code'),                        # `code`
            (r'\*\*\*([^*]+)\*\*\*', 'bold_italic'),      # ***text***
            (r'___([^_]+)___', 'bold_italic'),            # ___text___
            (r'\*\*([^*]+)\*\*', 'bold'),                 # **text**
            (r'__([^_]+)__', 'bold'),                     # __text__
            (r'\*([^*]+)\*', 'italic'),                   # *text*
            (r'_([^_]+)_', 'italic'),                     # _text_
        ]

        i = 0
        while i < len(text):
            matched = False
            for pattern, elem_type in patterns:
                match = re.match(pattern, text[i:], re.DOTALL)
                if match:
                    # Add text before match
                    if match.start() > 0:
                        parts.append(('text', text[i:i + match.start()]))

                    if elem_type == 'image':
                        parts.append(('image', match.group(1), match.group(2)))
                    elif elem_type == 'link':
                        parts.append(('link', match.group(1), match.group(2)))
                    elif elem_type == 'code':
                        parts.append(('code', match.group(1)))
                    elif elem_type == 'bold':
                        parts.append(('bold', match.group(1)))
                    elif elem_type == 'italic':
                        parts.append(('italic', match.group(1)))
                    elif elem_type == 'bold_italic':
                        parts.append(('bold_italic', match.group(1)))

                    i += match.end()
                    matched = True
                    break

            if not matched:
                # Find next potential match or end of text
                next_pos = len(text)
                for pattern, _ in patterns:
                    match = re.search(pattern, text[i:])
                    if match and match.start() < next_pos:
                        next_pos = match.start()

                if next_pos > 0:
                    parts.append(('text', text[i:i + next_pos]))
                    i += next_pos
                else:
                    break

        return parts

    def _add_formatted_paragraph(self, text):
        """Add a paragraph with inline formatting."""
        para = self.doc.add_paragraph()
        parts = self._parse_inline_formatting(text)

        for part in parts:
            if part[0] == 'text':
                run = para.add_run(part[1])
                self._set_run_font(run)
            elif part[0] == 'bold':
                run = para.add_run(part[1])
                self._set_run_font(run, bold=True)
            elif part[0] == 'italic':
                run = para.add_run(part[1])
                self._set_run_font(run, italic=True)
            elif part[0] == 'bold_italic':
                run = para.add_run(part[1])
                self._set_run_font(run, bold=True, italic=True)
            elif part[0] == 'code':
                run = para.add_run(part[1])
                self._set_run_font(run, code=True)
                # Add light gray background for code
                run.element.rPr  # Ensure rPr exists
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F5F5F5')
                run.element.rPr.append(shading_elm)
            elif part[0] == 'link':
                run = para.add_run(part[1])
                self._set_run_font(run)
                # Add hyperlink
                hyperlink = OxmlElement('w:hyperlink')
                hyperlink.set(qn('r:id'), part[2])
                run._element.append(hyperlink)
            elif part[0] == 'image':
                # Image handling - requires path or URL
                try:
                    # Check if it's a local file
                    img_path = Path(part[2])
                    if img_path.exists():
                        para.add_run().add_picture(str(img_path), width=Inches(5))
                    else:
                        run = para.add_run(f"[Image: {part[1]}]")
                        self._set_run_font(run, italic=True)
                except Exception:
                    run = para.add_run(f"[Image: {part[1]}]")
                    self._set_run_font(run, italic=True)

        return para

    def convert(self, md_text):
        """Convert Markdown text to Word document."""
        elements = self.parse_markdown_elements(md_text)

        for elem in elements:
            elem_type = elem[0]

            if elem_type == 'heading':
                level = min(elem[2], 9)  # Word has 9 heading levels
                heading = self.doc.add_heading(elem[1], level=level)
                # Ensure heading uses correct font
                for run in heading.runs:
                    self._set_run_font(run, bold=True)

            elif elem_type == 'paragraph':
                if elem[1].strip():
                    self._add_formatted_paragraph(elem[1])

            elif elem_type == 'ulist':
                for item in elem[1]:
                    para = self.doc.add_paragraph(item, style='List Bullet')
                    for run in para.runs:
                        self._set_run_font(run)

            elif elem_type == 'olist':
                for item in elem[1]:
                    para = self.doc.add_paragraph(item, style='List Number')
                    for run in para.runs:
                        self._set_run_font(run)

            elif elem_type == 'code':
                para = self.doc.add_paragraph()
                run = para.add_run(elem[1])
                self._set_run_font(run, code=True)
                # Add gray background
                run.element.rPr  # Ensure rPr exists
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F5F5F5')
                run.element.rPr.append(shading_elm)
                # Add border
                pPr = para._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                pPr.append(pBdr)
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '1')
                    border.set(qn('w:color'), 'D0D0D0')
                    pBdr.append(border)

            elif elem_type == 'table':
                headers = elem[1]
                rows = elem[2]
                num_cols = len(headers)

                # Create table
                table = self.doc.add_table(rows=len(rows) + 1, cols=num_cols)
                table.style = 'Light Grid Accent 1'

                # Add header row
                header_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    # Set header font
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            self._set_run_font(run, bold=True)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Add data rows
                for row_idx, row_data in enumerate(rows):
                    row_cells = table.rows[row_idx + 1].cells
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < num_cols:
                            row_cells[col_idx].text = cell_data
                            # Set cell font
                            for paragraph in row_cells[col_idx].paragraphs:
                                for run in paragraph.runs:
                                    self._set_run_font(run)

            elif elem_type == 'quote':
                para = self.doc.add_paragraph(elem[1])
                para.paragraph_format.left_indent = Inches(0.5)
                para.paragraph_format.right_indent = Inches(0.5)
                for run in para.runs:
                    self._set_run_font(run, italic=True)
                # Add left border
                pPr = para._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                pPr.append(pBdr)
                border = OxmlElement('w:left')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '6')
                border.set(qn('w:space'), '4')
                border.set(qn('w:color'), '4A90E2')
                pBdr.append(border)

            elif elem_type == 'hr':
                paragraph = self.doc.add_paragraph()
                pPr = paragraph._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                pPr.append(pBdr)
                border = OxmlElement('w:bottom')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '6')
                border.set(qn('w:space'), '1')
                border.set(qn('w:color'), 'CCCCCC')
                pBdr.append(border)

    def save(self, output_path):
        """Save the Word document."""
        self.doc.save(output_path)


def main():
    """Main entry point for CLI usage."""
    parser = argparse.ArgumentParser(description='Convert Markdown to Word document')
    parser.add_argument('input', help='Input Markdown file path')
    parser.add_argument('output', nargs='?', help='Output Word document path (.docx). Default: same directory as input with .docx extension')
    parser.add_argument('--chinese-font', default='宋体', help='Font for Chinese characters')
    parser.add_argument('--english-font', default='Arial', help='Font for English characters')
    parser.add_argument('--font-size', type=int, default=11, help='Base font size in points')

    args = parser.parse_args()

    # Read input file
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: Input file not found: {args.input}")
        sys.exit(1)

    md_text = input_path.read_text(encoding='utf-8')

    # Determine output path (default: same directory, same name, .docx extension)
    if args.output:
        output_path = Path(args.output)
    else:
        output_path = input_path.with_suffix('.docx')

    # Convert
    converter = MarkdownToWordConverter(
        chinese_font=args.chinese_font,
        english_font=args.english_font,
        font_size=args.font_size
    )
    converter.convert(md_text)
    converter.save(str(output_path))

    print(f"[OK] Converted: {input_path} -> {output_path}")


if __name__ == '__main__':
    main()
